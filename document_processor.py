"""
Offline Document Processor for Local LLM Context Injection
Supports: PDF, DOCX, XLSX, TXT, Images (OCR)
Target Model: Llama Nano 128K context window

No internet required - fully offline operation.
"""

import os
import hashlib
import logging
from pathlib import Path
from typing import List, Dict, Optional, Union
from dataclasses import dataclass, field
from enum import Enum
from datetime import datetime

# Extraction
import fitz  # pymupdf
from docx import Document as DocxDocument
import openpyxl
from PIL import Image
import pytesseract

# OCR Processing Pipeline (pre/post-processing for artifact handling)
try:
    from ocr_processor import OCRPipeline, preprocess_image, postprocess_text
    OCR_PIPELINE_AVAILABLE = True
except ImportError:
    OCR_PIPELINE_AVAILABLE = False

# Chunking & Embeddings
from sentence_transformers import SentenceTransformer

# Python 3.14 compatibility patch for ChromaDB
import chromadb_compat  # noqa: F401 - side-effect import for patching

import chromadb
from chromadb.config import Settings

logger = logging.getLogger(__name__)


class DocumentType(Enum):
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    XLSX = "xlsx"
    TXT = "txt"
    IMAGE = "image"
    UNKNOWN = "unknown"


@dataclass
class DocumentChunk:
    """Single chunk of extracted document content."""
    content: str
    metadata: Dict
    chunk_index: int
    doc_hash: str


@dataclass
class ProcessedDocument:
    """Metadata for a fully processed document."""
    doc_hash: str
    filename: str
    file_path: str
    doc_type: DocumentType
    chunks: List[DocumentChunk]
    total_tokens_estimate: int
    extracted_at: datetime = field(default_factory=datetime.now)
    ocr_used: bool = False
    ocr_pages: int = 0
    
    def to_dict(self) -> Dict:
        return {
            "doc_hash": self.doc_hash,
            "filename": self.filename,
            "file_path": self.file_path,
            "doc_type": self.doc_type.value,
            "chunk_count": len(self.chunks),
            "total_tokens_estimate": self.total_tokens_estimate,
            "extracted_at": self.extracted_at.isoformat(),
            "ocr_used": self.ocr_used,
            "ocr_pages": self.ocr_pages
        }


class DocumentExtractor:
    """Handles raw text extraction from various document formats."""
    
    SUPPORTED_IMAGES = {'.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif', '.webp'}
    SUPPORTED_DOCS = {'.pdf', '.docx', '.doc', '.xlsx', '.xls', '.txt', '.md', '.csv'}
    
    def __init__(self, tesseract_path: Optional[str] = None):
        """
        Initialize extractor.
        
        Args:
            tesseract_path: Path to tesseract executable (e.g., 
                           "C:/Program Files/Tesseract-OCR/tesseract.exe" on Windows)
        """
        if tesseract_path:
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
        
        self._tesseract_available = self._check_tesseract()
    
    def _check_tesseract(self) -> bool:
        """Verify tesseract is available."""
        try:
            pytesseract.get_tesseract_version()
            return True
        except Exception as e:
            logger.warning(f"Tesseract not available: {e}. OCR disabled.")
            return False
    
    @classmethod
    def get_supported_extensions(cls) -> set:
        """Return all supported file extensions."""
        return cls.SUPPORTED_IMAGES | cls.SUPPORTED_DOCS
    
    def extract(self, file_path: Path) -> str:
        """
        Extract text from document.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Tuple of (extracted_text, ocr_used, ocr_page_count)
            
        Raises:
            ValueError: If file type not supported
            FileNotFoundError: If file doesn't exist
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        suffix = file_path.suffix.lower()
        
        # Methods that return OCR metadata
        ocr_extractors = {
            '.pdf': self._extract_pdf,
        }
        
        # Methods that don't use OCR (return just text)
        text_extractors = {
            '.docx': self._extract_docx,
            '.doc': self._extract_doc_legacy,
            '.xlsx': self._extract_xlsx,
            '.xls': self._extract_xlsx,
            '.txt': self._extract_txt,
            '.md': self._extract_txt,
            '.csv': self._extract_csv,
        }
        
        if suffix in ocr_extractors:
            return ocr_extractors[suffix](file_path)
        elif suffix in text_extractors:
            # Wrap result in tuple with OCR=False
            text = text_extractors[suffix](file_path)
            return text, False, 0
        elif suffix in self.SUPPORTED_IMAGES:
            return self._extract_image_ocr(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
    
    def _extract_pdf(self, path: Path) -> tuple[str, bool, int]:
        """
        Extract text from PDF with OCR fallback for scanned pages.

        Uses two-step OCR processing:
        1. Pre-processing: Image enhancement (contrast, denoise, deskew)
        2. Post-processing: Text cleanup (fix artifacts, normalize whitespace)

        Returns:
            Tuple of (extracted_text, ocr_was_used, ocr_page_count)
        """
        doc = fitz.open(path)
        text_parts = []
        ocr_used = False
        ocr_page_count = 0

        # Initialize OCR pipeline if available
        ocr_pipeline = None
        if OCR_PIPELINE_AVAILABLE and self._tesseract_available:
            ocr_pipeline = OCRPipeline()

        for page_num, page in enumerate(doc):
            text = page.get_text()

            # If page has minimal text, attempt OCR with pre/post processing
            if len(text.strip()) < 50 and self._tesseract_available:
                try:
                    # Render page to image at higher DPI for better OCR
                    pix = page.get_pixmap(dpi=200)
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

                    if ocr_pipeline:
                        # Full pipeline: preprocess -> OCR -> postprocess
                        ocr_text, metadata = ocr_pipeline.process_image(img)
                        logger.debug(
                            f"OCR page {page_num + 1}: "
                            f"{metadata['raw_length']} -> {metadata['clean_length']} chars "
                            f"({metadata['reduction_pct']}% reduction)"
                        )
                    else:
                        # Fallback: basic OCR without pre/post processing
                        ocr_text = pytesseract.image_to_string(img)

                    if ocr_text.strip():
                        text = f"[OCR]\n{ocr_text}"
                        ocr_used = True
                        ocr_page_count += 1
                        logger.debug(f"OCR applied to page {page_num + 1}")
                except Exception as e:
                    logger.warning(f"OCR failed for page {page_num + 1}: {e}")
            elif text.strip() and OCR_PIPELINE_AVAILABLE:
                # Even for native text, apply post-processing to clean up
                text = postprocess_text(text)

            if text.strip():
                text_parts.append(f"[Page {page_num + 1}]\n{text.strip()}")

        doc.close()
        return "\n\n".join(text_parts), ocr_used, ocr_page_count
    
    def _extract_docx(self, path: Path) -> str:
        """Extract text from DOCX."""
        doc = DocxDocument(path)
        paragraphs = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)
        
        return "\n\n".join(paragraphs)
    
    def _extract_doc_legacy(self, path: Path) -> str:
        """
        Handle legacy .doc format.
        
        Note: Full support requires antiword or LibreOffice.
        """
        # Try using textract if available
        try:
            import textract
            return textract.process(str(path)).decode('utf-8')
        except ImportError:
            pass
        
        # Fallback: inform user
        raise NotImplementedError(
            f"Legacy .doc format requires conversion. "
            f"Please save '{path.name}' as .docx, or install 'textract' package."
        )
    
    def _extract_xlsx(self, path: Path) -> str:
        """Extract text from Excel spreadsheet."""
        wb = openpyxl.load_workbook(path, data_only=True)
        sheets_text = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            rows = []
            
            for row in sheet.iter_rows(values_only=True):
                # Filter None values and convert to string
                row_values = [str(cell) if cell is not None else "" for cell in row]
                row_text = " | ".join(row_values)
                
                # Skip empty rows
                if row_text.strip(" |"):
                    rows.append(row_text)
            
            if rows:
                sheets_text.append(f"[Sheet: {sheet_name}]\n" + "\n".join(rows))
        
        wb.close()
        return "\n\n".join(sheets_text)
    
    def _extract_txt(self, path: Path) -> str:
        """Extract text from plain text file."""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                return path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        
        # Last resort: ignore errors
        return path.read_text(encoding='utf-8', errors='replace')
    
    def _extract_csv(self, path: Path) -> str:
        """Extract text from CSV file."""
        import csv
        
        text = self._extract_txt(path)
        lines = text.strip().split('\n')
        
        if not lines:
            return ""
        
        # Format as table
        formatted = []
        for line in lines[:1000]:  # Limit to first 1000 rows
            formatted.append(line.replace(',', ' | '))
        
        return "\n".join(formatted)
    
    def _extract_image_ocr(self, path: Path) -> tuple[str, bool, int]:
        """
        Extract text from image using OCR with pre/post processing.

        Uses two-step processing:
        1. Pre-processing: Image enhancement (contrast, denoise, threshold)
        2. Post-processing: Text cleanup (fix artifacts, common errors)

        Returns:
            Tuple of (extracted_text, ocr_was_used, ocr_page_count)
        """
        if not self._tesseract_available:
            raise RuntimeError(
                "Tesseract OCR not available. Please install Tesseract: "
                "https://github.com/tesseract-ocr/tesseract"
            )

        img = Image.open(path)

        # Convert to RGB if necessary
        if img.mode != 'RGB':
            img = img.convert('RGB')

        if OCR_PIPELINE_AVAILABLE:
            # Full pipeline: preprocess -> OCR -> postprocess
            ocr_pipeline = OCRPipeline()
            text, metadata = ocr_pipeline.process_image(img)
            logger.info(
                f"OCR image {path.name}: "
                f"{metadata['raw_length']} -> {metadata['clean_length']} chars "
                f"({metadata['reduction_pct']}% artifact reduction)"
            )
        else:
            # Fallback: basic OCR
            text = pytesseract.image_to_string(img)

        img.close()

        # Images always use OCR, count as 1 page
        return text, True, 1


class SemanticChunker:
    """
    Token-aware semantic chunking optimized for large context windows.
    
    Designed for Llama Nano 128K - chunks are sized to allow multiple
    relevant chunks while preserving context.
    """
    
    def __init__(
        self,
        chunk_size: int = 1000,       # target tokens per chunk
        chunk_overlap: int = 100,      # overlap tokens for continuity
        chars_per_token: float = 4.0   # rough estimate (varies by model)
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.chars_per_token = chars_per_token
    
    def chunk(self, text: str, doc_hash: str) -> List[DocumentChunk]:
        """
        Split text into overlapping semantic chunks.
        
        Args:
            text: Full document text
            doc_hash: Document identifier
            
        Returns:
            List of DocumentChunk objects
        """
        if not text.strip():
            return []
        
        char_chunk_size = int(self.chunk_size * self.chars_per_token)
        char_overlap = int(self.chunk_overlap * self.chars_per_token)
        
        chunks = []
        start = 0
        chunk_index = 0
        
        while start < len(text):
            end = start + char_chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                boundary = self._find_sentence_boundary(text, end)
                if boundary > start + (char_chunk_size // 2):  # Ensure minimum chunk size
                    end = boundary
            
            chunk_text = text[start:end].strip()
            
            if chunk_text:
                chunks.append(DocumentChunk(
                    content=chunk_text,
                    metadata={
                        "start_char": start,
                        "end_char": end,
                        "char_count": len(chunk_text)
                    },
                    chunk_index=chunk_index,
                    doc_hash=doc_hash
                ))
                chunk_index += 1
            
            # Move forward with overlap
            start = end - char_overlap
            if start <= chunks[-1].metadata["start_char"] if chunks else 0:
                start = end  # Prevent infinite loop
        
        return chunks
    
    def _find_sentence_boundary(self, text: str, pos: int, window: int = 200) -> int:
        """Find nearest sentence boundary near position."""
        search_start = max(0, pos - window)
        search_end = min(len(text), pos + window)
        
        # Priority order for sentence boundaries
        boundaries = ['. ', '.\n', '.\t', '? ', '!\n', '?\n', '! ', '\n\n']
        
        best_pos = pos
        best_distance = window
        
        for boundary in boundaries:
            idx = text.rfind(boundary, search_start, pos)
            if idx > search_start:
                distance = pos - idx
                if distance < best_distance:
                    best_pos = idx + len(boundary)
                    best_distance = distance
        
        return best_pos


class LocalVectorStore:
    """
    ChromaDB-backed local vector store for semantic search.
    
    Fully offline - uses local sentence-transformers model.
    """
    
    def __init__(
        self,
        persist_directory: Optional[str] = None,
        embedding_model: str = "all-MiniLM-L6-v2",
        collection_name: str = "jan_documents"
    ):
        """
        Initialize vector store.
        
        Args:
            persist_directory: Path for persistent storage (None for ephemeral)
            embedding_model: sentence-transformers model name
            collection_name: ChromaDB collection name
        """
        logger.info(f"Loading embedding model: {embedding_model}")
        self.embedder = SentenceTransformer(embedding_model)
        
        settings = Settings(
            anonymized_telemetry=False,
            allow_reset=True
        )
        
        if persist_directory:
            os.makedirs(persist_directory, exist_ok=True)
            self.client = chromadb.PersistentClient(
                path=persist_directory,
                settings=settings
            )
            logger.info(f"Using persistent storage: {persist_directory}")
        else:
            self.client = chromadb.EphemeralClient(settings=settings)
            logger.info("Using ephemeral storage")
        
        self.collection = self.client.get_or_create_collection(
            name=collection_name,
            metadata={"hnsw:space": "cosine"}
        )
    
    def add_chunks(self, chunks: List[DocumentChunk], filename: str):
        """Add document chunks to vector store."""
        if not chunks:
            return
        
        ids = [f"{chunks[0].doc_hash}_{c.chunk_index}" for c in chunks]
        documents = [c.content for c in chunks]
        
        logger.info(f"Embedding {len(chunks)} chunks...")
        embeddings = self.embedder.encode(documents, show_progress_bar=True).tolist()
        
        metadatas = [
            {
                **c.metadata,
                "filename": filename,
                "doc_hash": c.doc_hash,
                "chunk_index": c.chunk_index
            }
            for c in chunks
        ]
        
        self.collection.add(
            ids=ids,
            documents=documents,
            embeddings=embeddings,
            metadatas=metadatas
        )
        
        logger.info(f"Added {len(chunks)} chunks from {filename}")
    
    def query(
        self,
        query_text: str,
        n_results: int = 5,
        filter_doc_hash: Optional[str] = None
    ) -> List[Dict]:
        """
        Query for similar chunks.
        
        Args:
            query_text: Search query
            n_results: Number of results to return
            filter_doc_hash: Optionally filter to specific document
            
        Returns:
            List of result dicts with content, metadata, distance
        """
        query_embedding = self.embedder.encode([query_text]).tolist()
        
        where_filter = {"doc_hash": filter_doc_hash} if filter_doc_hash else None
        
        results = self.collection.query(
            query_embeddings=query_embedding,
            n_results=n_results,
            where=where_filter,
            include=["documents", "metadatas", "distances"]
        )
        
        if not results["documents"] or not results["documents"][0]:
            return []
        
        return [
            {
                "content": doc,
                "metadata": meta,
                "distance": dist,
                "relevance_score": 1 - dist  # Convert distance to similarity
            }
            for doc, meta, dist in zip(
                results["documents"][0],
                results["metadatas"][0],
                results["distances"][0]
            )
        ]
    
    def delete_document(self, doc_hash: str):
        """Delete all chunks for a document."""
        self.collection.delete(where={"doc_hash": doc_hash})
        logger.info(f"Deleted document: {doc_hash}")
    
    def get_document_count(self) -> int:
        """Get total number of chunks in store."""
        return self.collection.count()
    
    def list_documents(self) -> List[Dict]:
        """List all unique documents in store."""
        # Get all metadata
        results = self.collection.get(include=["metadatas"])
        
        if not results["metadatas"]:
            return []
        
        # Deduplicate by doc_hash
        docs = {}
        for meta in results["metadatas"]:
            doc_hash = meta.get("doc_hash")
            if doc_hash and doc_hash not in docs:
                docs[doc_hash] = {
                    "doc_hash": doc_hash,
                    "filename": meta.get("filename", "unknown")
                }
        
        return list(docs.values())


class DocumentProcessor:
    """
    Main interface for offline document processing.
    
    Orchestrates extraction, chunking, and vector storage.
    
    Usage:
        processor = DocumentProcessor(persist_directory="./doc_store")
        processor.ingest("/path/to/document.pdf")
        context = processor.get_context("What does this say about X?")
    """
    
    def __init__(
        self,
        persist_directory: Optional[str] = None,
        tesseract_path: Optional[str] = None,
        embedding_model: str = "all-MiniLM-L6-v2",
        chunk_size: int = 1000,
        chunk_overlap: int = 100
    ):
        """
        Initialize document processor.
        
        Args:
            persist_directory: Path for persistent vector storage
            tesseract_path: Path to tesseract executable for OCR
            embedding_model: sentence-transformers model for embeddings
            chunk_size: Target tokens per chunk
            chunk_overlap: Overlap tokens between chunks
        """
        self.extractor = DocumentExtractor(tesseract_path=tesseract_path)
        self.chunker = SemanticChunker(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
        self.vector_store = LocalVectorStore(
            persist_directory=persist_directory,
            embedding_model=embedding_model
        )
        self.processed_docs: Dict[str, ProcessedDocument] = {}
    
    def _compute_hash(self, file_path: Path) -> str:
        """Compute unique hash for file content."""
        hasher = hashlib.sha256()
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(8192), b''):
                hasher.update(chunk)
        return hasher.hexdigest()[:16]
    
    def _detect_type(self, path: Path) -> DocumentType:
        """Detect document type from extension."""
        suffix = path.suffix.lower()
        
        type_map = {
            '.pdf': DocumentType.PDF,
            '.docx': DocumentType.DOCX,
            '.doc': DocumentType.DOC,
            '.xlsx': DocumentType.XLSX,
            '.xls': DocumentType.XLSX,
            '.txt': DocumentType.TXT,
            '.md': DocumentType.TXT,
            '.csv': DocumentType.TXT,
        }
        
        if suffix in type_map:
            return type_map[suffix]
        elif suffix in DocumentExtractor.SUPPORTED_IMAGES:
            return DocumentType.IMAGE
        else:
            return DocumentType.UNKNOWN
    
    def ingest(self, file_path: Union[str, Path], force: bool = False) -> ProcessedDocument:
        """
        Ingest and index a document.
        
        Args:
            file_path: Path to document
            force: Re-process even if already indexed
            
        Returns:
            ProcessedDocument with metadata including OCR info
        """
        path = Path(file_path).resolve()
        
        if not path.exists():
            raise FileNotFoundError(f"Document not found: {path}")
        
        doc_hash = self._compute_hash(path)
        
        # Skip if already processed (unless forced)
        if not force and doc_hash in self.processed_docs:
            logger.info(f"Document already indexed: {path.name}")
            return self.processed_docs[doc_hash]
        
        logger.info(f"Processing: {path.name}")
        
        # Extract text (now returns OCR metadata)
        raw_text, ocr_used, ocr_pages = self.extractor.extract(path)
        
        if ocr_used:
            logger.info(f"OCR applied to {path.name}: {ocr_pages} page(s)")
        
        if not raw_text.strip():
            logger.warning(f"No text extracted from: {path.name}")
        
        # Chunk
        chunks = self.chunker.chunk(raw_text, doc_hash)
        
        # Create record with OCR metadata
        processed = ProcessedDocument(
            doc_hash=doc_hash,
            filename=path.name,
            file_path=str(path),
            doc_type=self._detect_type(path),
            chunks=chunks,
            total_tokens_estimate=int(len(raw_text) / 4),
            ocr_used=ocr_used,
            ocr_pages=ocr_pages
        )
        
        # Store in vector DB
        if chunks:
            self.vector_store.add_chunks(chunks, path.name)
        
        self.processed_docs[doc_hash] = processed
        
        ocr_info = f", OCR: {ocr_pages} pages" if ocr_used else ""
        logger.info(
            f"Indexed {path.name}: {len(chunks)} chunks, "
            f"~{processed.total_tokens_estimate} tokens{ocr_info}"
        )
        
        return processed
    
    def ingest_directory(
        self,
        directory: Union[str, Path],
        recursive: bool = True,
        extensions: Optional[set] = None
    ) -> List[ProcessedDocument]:
        """
        Ingest all supported documents in a directory.
        
        Args:
            directory: Directory path
            recursive: Include subdirectories
            extensions: Filter to specific extensions (None = all supported)
            
        Returns:
            List of ProcessedDocument objects
        """
        dir_path = Path(directory)
        if not dir_path.is_dir():
            raise NotADirectoryError(f"Not a directory: {directory}")
        
        supported = extensions or DocumentExtractor.get_supported_extensions()
        results = []
        
        pattern = "**/*" if recursive else "*"
        
        for file_path in dir_path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in supported:
                try:
                    result = self.ingest(file_path)
                    results.append(result)
                except Exception as e:
                    logger.error(f"Failed to process {file_path}: {e}")
        
        return results
    
    def get_context(
        self,
        query: str,
        n_chunks: int = 5,
        max_tokens: int = 8000,
        doc_hash: Optional[str] = None
    ) -> str:
        """
        Retrieve relevant context for LLM prompt injection.
        
        Args:
            query: User's question
            n_chunks: Max chunks to retrieve
            max_tokens: Token budget for context
            doc_hash: Optionally filter to specific document
            
        Returns:
            Formatted context string ready for prompt injection
        """
        results = self.vector_store.query(
            query,
            n_results=n_chunks,
            filter_doc_hash=doc_hash
        )
        
        if not results:
            return ""
        
        context_parts = []
        token_estimate = 0
        
        for r in results:
            chunk_tokens = len(r["content"]) / 4
            
            if token_estimate + chunk_tokens > max_tokens:
                break
            
            # Format with source attribution
            source = r["metadata"].get("filename", "unknown")
            relevance = r.get("relevance_score", 0)
            
            context_parts.append(
                f"[Source: {source} | Relevance: {relevance:.2f}]\n{r['content']}"
            )
            token_estimate += chunk_tokens
        
        return "\n\n---\n\n".join(context_parts)
    
    def remove_document(self, file_path: Union[str, Path]):
        """Remove a document from the index by file path."""
        path = Path(file_path).resolve()
        doc_hash = self._compute_hash(path)
        self.remove_document_by_hash(doc_hash)

    def remove_document_by_hash(self, doc_hash: str):
        """Remove a document from the index by its hash."""
        if doc_hash in self.processed_docs:
            filename = self.processed_docs[doc_hash].filename
            self.vector_store.delete_document(doc_hash)
            del self.processed_docs[doc_hash]
            logger.info(f"Removed: {filename} ({doc_hash})")
    
    def list_documents(self) -> List[Dict]:
        """List all indexed documents."""
        return [doc.to_dict() for doc in self.processed_docs.values()]
    
    def get_stats(self) -> Dict:
        """Get processor statistics."""
        return {
            "documents_indexed": len(self.processed_docs),
            "total_chunks": self.vector_store.get_document_count(),
            "supported_extensions": list(DocumentExtractor.get_supported_extensions())
        }

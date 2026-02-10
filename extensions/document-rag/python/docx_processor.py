#!/usr/bin/env python3
"""
DOCX Processor for AVACHATTER
Handles Microsoft Word document text extraction.
"""

from docx import Document
from typing import Dict, List
import os


class DOCXProcessor:
    """Extract text from Microsoft Word documents."""

    def extract_text(self, docx_path: str, preserve_formatting: bool = False) -> Dict[str, any]:
        """
        Extract text from DOCX file.

        Args:
            docx_path: Path to DOCX file
            preserve_formatting: Whether to preserve paragraph structure

        Returns:
            Dictionary with:
                - text: Extracted text
                - paragraphs: Number of paragraphs
                - tables: Number of tables
                - images: Number of images
                - metadata: Document properties
        """
        result = {
            'text': '',
            'paragraphs': 0,
            'tables': 0,
            'images': 0,
            'metadata': {},
            'error': None
        }

        try:
            doc = Document(docx_path)

            # Extract metadata
            result['metadata'] = self._extract_metadata(doc)

            # Extract text from paragraphs
            text_chunks = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_chunks.append(para.text)

            result['paragraphs'] = len(doc.paragraphs)

            # Extract text from tables
            for table in doc.tables:
                table_text = self._extract_table_text(table)
                if table_text:
                    text_chunks.append(table_text)

            result['tables'] = len(doc.tables)

            # Count images
            result['images'] = self._count_images(doc)

            # Join text
            if preserve_formatting:
                result['text'] = '\n\n'.join(text_chunks)
            else:
                result['text'] = ' '.join(text_chunks)

        except Exception as e:
            result['error'] = str(e)

        return result

    def _extract_metadata(self, doc: Document) -> Dict[str, str]:
        """
        Extract document metadata/properties.

        Args:
            doc: Document object

        Returns:
            Dictionary of metadata
        """
        metadata = {}

        try:
            core_props = doc.core_properties
            if core_props.author:
                metadata['author'] = core_props.author
            if core_props.title:
                metadata['title'] = core_props.title
            if core_props.subject:
                metadata['subject'] = core_props.subject
            if core_props.created:
                metadata['created'] = str(core_props.created)
            if core_props.modified:
                metadata['modified'] = str(core_props.modified)
        except:
            pass

        return metadata

    def _extract_table_text(self, table) -> str:
        """
        Extract text from a table.

        Args:
            table: Table object

        Returns:
            Text representation of table
        """
        table_text = []

        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)

            if row_text:
                table_text.append(' | '.join(row_text))

        return '\n'.join(table_text)

    def _count_images(self, doc: Document) -> int:
        """
        Count number of images in document.

        Args:
            doc: Document object

        Returns:
            Number of images
        """
        try:
            return len(doc.inline_shapes)
        except:
            return 0

    def extract_paragraphs(self, docx_path: str) -> List[str]:
        """
        Extract paragraphs as a list.

        Args:
            docx_path: Path to DOCX file

        Returns:
            List of paragraph texts
        """
        try:
            doc = Document(docx_path)
            return [para.text for para in doc.paragraphs if para.text.strip()]
        except Exception as e:
            return [f"Error: {str(e)}"]

    def extract_headings(self, docx_path: str) -> List[Dict[str, str]]:
        """
        Extract document headings with levels.

        Args:
            docx_path: Path to DOCX file

        Returns:
            List of dictionaries with 'level' and 'text'
        """
        headings = []

        try:
            doc = Document(docx_path)

            for para in doc.paragraphs:
                if para.style.name.startswith('Heading'):
                    level = para.style.name.replace('Heading ', '')
                    headings.append({
                        'level': level,
                        'text': para.text
                    })
        except Exception as e:
            headings.append({'level': 'error', 'text': str(e)})

        return headings


def main():
    """Test the DOCX processor."""
    import sys

    if len(sys.argv) < 2:
        print("Usage: python docx_processor.py <docx_file>")
        sys.exit(1)

    docx_path = sys.argv[1]

    if not os.path.exists(docx_path):
        print(f"File not found: {docx_path}")
        sys.exit(1)

    processor = DOCXProcessor()

    # Extract text
    print("Extracting text...")
    result = processor.extract_text(docx_path, preserve_formatting=True)

    if result['error']:
        print(f"Error: {result['error']}")
        sys.exit(1)

    print(f"\nParagraphs: {result['paragraphs']}")
    print(f"Tables: {result['tables']}")
    print(f"Images: {result['images']}")
    print(f"\nMetadata: {result['metadata']}")
    print(f"\nExtracted Text ({len(result['text'])} characters):")
    print("=" * 80)
    print(result['text'][:1000])  # Print first 1000 chars
    if len(result['text']) > 1000:
        print("...")

    # Extract headings
    print("\n\nDocument Structure (Headings):")
    print("=" * 80)
    headings = processor.extract_headings(docx_path)
    for heading in headings[:10]:  # Show first 10
        indent = "  " * (int(heading['level']) if heading['level'].isdigit() else 0)
        print(f"{indent}{heading['text']}")


if __name__ == '__main__':
    main()
